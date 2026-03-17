from docx import Document
import re

def convert_word(text, topic):
    doc = Document()

    lines = text.split("\n")

    for line in lines:
        line = line.strip()

        if not line:
            continue

        if line.startswith("Title:"):
            doc.add_heading(line.replace("Title:", "").strip(), level=0)

        elif line.startswith("Abstract"):
            doc.add_heading("Abstract", level=1)

        elif re.match(r"^\d+\s+[A-Za-z]", line):
            doc.add_heading(line, level=1)

        elif re.match(r"^\d+\.\d+\s", line):
            doc.add_heading(line, level=2)

        elif re.match(r"^\d+\.\d+\.\d+\s", line):
            doc.add_heading(line, level=3)

        elif line.lower().startswith("references"):
            doc.add_heading("References", level=1)

        elif re.match(r"^\[\d+\]", line):
            doc.add_paragraph(line)

        else:
            doc.add_paragraph(line)

    doc.save(f"{topic[:15]}_review_paper.docx")